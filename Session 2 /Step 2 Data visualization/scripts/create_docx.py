#!/usr/bin/env python3
"""
Create a WHU-styled DOCX report from a JSON spec.

Usage:
  python scripts/create_docx.py --spec assets/example_docx_spec.json --out out.docx
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WHU_DARK_BLUE = RGBColor(0x05, 0x46, 0x96)
WHU_LIGHT_BLUE = RGBColor(0x00, 0xA5, 0xDC)

def _set_base_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(20)
    h1.font.color.rgb = WHU_DARK_BLUE

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.color.rgb = WHU_LIGHT_BLUE

def build(spec_path: Path, out_path: Path):
    spec = json.loads(spec_path.read_text(encoding="utf-8"))
    doc = Document()
    _set_base_styles(doc)

    # Cover
    title = spec.get("title", "")
    subtitle = spec.get("subtitle", "")
    author_line = spec.get("author_line", "")

    p = doc.add_paragraph(title, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if subtitle:
        doc.add_paragraph(subtitle, style="Heading 2")
    if author_line:
        ap = doc.add_paragraph(author_line)
        ap.runs[0].font.name = "Arial"
        ap.runs[0].font.size = Pt(11)
        ap.runs[0].font.color.rgb = WHU_DARK_BLUE

    doc.add_paragraph("")  # spacer

    for sec in spec.get("sections", []):
        heading = sec.get("heading")
        if heading:
            doc.add_paragraph(heading, style="Heading 1")
        for para in sec.get("paragraphs", []):
            doc.add_paragraph(para)
        bullets = sec.get("bullets") or []
        for b in bullets:
            doc.add_paragraph(str(b), style="List Bullet")

    doc.save(str(out_path))

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--spec", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    args = ap.parse_args()
    build(args.spec, args.out)

if __name__ == "__main__":
    main()
